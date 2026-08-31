"""
Shared helpers for creating/updating one-pagers/<date>/00_summary.docx.

Always ADDS -- never overwrites or deletes -- consistent with the routine's
add-not-overwrite rule for a folder that already has content from an earlier
run/batch the same day. See onepager_lib.py's module docstring for why this
avoids python-docx's by-name style lookup.
"""
import os
from docx import Document
from docx.shared import Pt

from onepager_lib import heading1, heading2, body, add_borders, shaded_header_row, data_row

KEEP_HEADERS = ["Screen", "Ticker", "Market Cap", "Raise Size", "% of Mkt Cap", "Urgency"]
DISPOSITION_HEADERS = ["Screen", "Ticker", "Score", "Disposition", "Reason"]

SUMMARY_FILENAME = "00_summary.docx"


def load_or_create(folder, week_of):
    path = os.path.join(folder, SUMMARY_FILENAME)
    if os.path.exists(path):
        doc = Document(path)
        keep_table = doc.tables[0] if doc.tables else None
        return doc, path, keep_table, True
    os.makedirs(folder, exist_ok=True)
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Weekly Equity-Raise Candidate Research — Summary")
    r.bold = True
    r.font.size = Pt(20)
    body(doc, f"Week of {week_of}")
    heading1(doc, "Keep Names — One-Pager Summary (Both Screens, Accumulated)")
    body(doc, "Every Keep name with a finished one-pager in this folder, across all batches researched this week.")
    keep_table = doc.add_table(rows=1, cols=len(KEEP_HEADERS))
    shaded_header_row(keep_table, KEEP_HEADERS)
    add_borders(keep_table)
    return doc, path, keep_table, False


def append_keep_rows(keep_table, rows):
    """rows: list of (screen, ticker, market_cap, raise_size, pct, urgency)"""
    for row in rows:
        data_row(keep_table, row)


def append_disposition_section(doc, batch_label, rows):
    """rows: list of (screen, ticker, score, disposition, reason)"""
    heading1(doc, f"Full Disposition — {batch_label}")
    t = doc.add_table(rows=1, cols=len(DISPOSITION_HEADERS))
    shaded_header_row(t, DISPOSITION_HEADERS, fill="2E74B5")
    for row in rows:
        data_row(t, row, size=8)
    add_borders(t)
    return t


def save(doc, path):
    doc.save(path)
    return path
